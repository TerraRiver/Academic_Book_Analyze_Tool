import os
import json
import time
import configparser
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.section import WD_SECTION_START
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
from typing import List, Dict, Optional, Callable
from .pdf_processor import get_leaf_chapters, enrich_chapters

PARENT_CHAPTER_ANALYSIS_PREFIX = "parent_"

class ReportGenerator:
    """
    负责生成Word报告的类。
    将LLM分析结果（Markdown文件）整合到Word文档中。
    """
    def __init__(self, book_path: str, log_callback: Optional[Callable] = None):
        self.book_path = book_path
        self.llm_result_dir = os.path.join(book_path, "LLM_result")
        self.log_callback = log_callback
        self.document = Document()
        self.config = configparser.ConfigParser()

        # 读取配置文件以确定输出目录
        self.config.read('config.ini', encoding='utf-8')
        custom_output_path = self.config.get('General', 'report_output_path', fallback=None)
        
        if custom_output_path and os.path.isdir(custom_output_path):
            self.output_dir = custom_output_path
            self._log(f"报告将保存到自定义目录: {self.output_dir}")
        else:
            self.output_dir = book_path # 默认输出到书籍根目录
        
        self._setup_document_defaults()

    def _log(self, message: str):
        if self.log_callback:
            self.log_callback(message)
        else:
            print(message)

    def _setup_document_defaults(self):
        """设置文档的默认样式和布局"""
        # 设置页边距
        section = self.document.sections[0]
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

        # 设置默认字体
        style = self.document.styles['Normal']
        font = style.font
        font.name = '宋体' # 可以根据需要调整
        font.size = Pt(12)
        # 设置中文字体
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def generate_report(self, book_title: str, chapters: List[Dict]) -> Optional[str]:
        """
        生成Word报告。
        
        Args:
            book_title: 书籍的标题。
            chapters: 章节信息列表，包含'title', 'start_page', 'end_page'等。
            
        Returns:
            生成的报告文件路径，如果失败则返回None。
        """
        self._log(f"开始生成书籍 '{book_title}' 的Word报告...")

        # 添加标题页
        self.document.add_heading(book_title, level=0)
        self.document.add_paragraph("LLM分析报告", style='Subtitle')
        self.document.add_paragraph(f"生成日期: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        self.document.add_page_break()

        # 检查LLM结果目录是否存在
        if not os.path.exists(self.llm_result_dir):
            self._log(f"LLM结果目录不存在: {self.llm_result_dir}")
            return None
        
        enriched_chapters = enrich_chapters(chapters)
        leaf_chapters = [chapter for chapter in enriched_chapters if chapter.get("is_leaf")]
        chapter_files = sorted([f for f in os.listdir(self.llm_result_dir) if f.endswith('_analysis.md')])

        if not chapter_files or not leaf_chapters:
            self._log("未找到任何LLM分析结果文件，无法生成报告。")
            return None

        leaf_analysis_map = self._load_leaf_analysis_map(leaf_chapters)
        enable_parent_summary_analysis = self.config.getboolean(
            'LLM', 'enable_parent_summary_analysis', fallback=True
        )
        parent_analysis_map = (
            self._load_parent_analysis_map(enriched_chapters)
            if enable_parent_summary_analysis else {}
        )
        leaf_chapter_keys = {self._chapter_key(chapter) for chapter in leaf_chapters}
        first_top_level = True

        for chapter in enriched_chapters:
            level = max(1, int(chapter.get("level", 1)))
            heading_text = chapter.get("title", "未命名章节")

            if level == 1 and not first_top_level:
                self.document.add_page_break()
            if level == 1:
                first_top_level = False

            self.document.add_heading(heading_text, level=min(level, 4))

            parent_analysis = parent_analysis_map.get(self._chapter_key(chapter))
            if parent_analysis:
                self._add_markdown_content_to_document(
                    parent_analysis,
                    skip_title_heading=True
                )

            if self._chapter_key(chapter) not in leaf_chapter_keys:
                continue

            analysis_content = leaf_analysis_map.get(self._chapter_key(chapter))
            if not analysis_content:
                self._log(f"警告: 未找到叶子章节分析结果: {heading_text}")
                self.document.add_paragraph("未找到对应的 LLM 分析结果。")
                continue

            self._add_markdown_content_to_document(
                analysis_content,
                skip_title_heading=True
            )
            self._log(f"已添加切片分析: {heading_text}")

        output_filename = f"{book_title}_Report.docx"
        output_path = os.path.join(self.output_dir, output_filename)

        try:
            self.document.save(output_path)
            self._log(f"Word报告生成成功: {output_path}")
            return output_path
        except Exception as e:
            self._log(f"保存Word报告失败: {e}")
            return None

    def _chapter_key(self, chapter: Dict) -> tuple:
        return (
            int(chapter.get("level", 1)),
            chapter.get("title", ""),
            int(chapter.get("start_page", 0)),
            int(chapter.get("end_page", 0)),
        )

    def _load_leaf_analysis_map(self, leaf_chapters: List[Dict]) -> Dict[tuple, str]:
        analysis_map = {}
        for i, chapter in enumerate(leaf_chapters):
            chapter_file = f"{i+1:02d}_analysis.md"
            file_path = os.path.join(self.llm_result_dir, chapter_file)
            if not os.path.exists(file_path):
                self._log(f"警告: 分析结果文件不存在: {chapter_file}")
                continue
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    analysis_map[self._chapter_key(chapter)] = f.read()
            except Exception as e:
                self._log(f"读取分析结果文件 '{chapter_file}' 失败: {e}")
        return analysis_map

    def _load_parent_analysis_map(self, chapters: List[Dict]) -> Dict[tuple, str]:
        analysis_map = {}
        for chapter in chapters:
            if chapter.get("is_leaf"):
                continue
            chapter_file = (
                f"{PARENT_CHAPTER_ANALYSIS_PREFIX}"
                f"{chapter.get('_source_index', 0) + 1:02d}_analysis.md"
            )
            file_path = os.path.join(self.llm_result_dir, chapter_file)
            if not os.path.exists(file_path):
                continue
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    analysis_map[self._chapter_key(chapter)] = f.read()
            except Exception as e:
                self._log(f"读取分析结果文件 '{chapter_file}' 失败: {e}")
        return analysis_map

    def _add_markdown_content_to_document(self, markdown_content: str, skip_title_heading: bool = False):
        """
        将Markdown内容解析并添加到Word文档中。
        这是一个简化的Markdown解析，处理标题、段落、粗体文本、分隔线等。
        """
        lines = markdown_content.split('\n')
        i = 0
        title_skipped = not skip_title_heading
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue

            if line.startswith('#'):
                if skip_title_heading and not title_skipped:
                    title_skipped = True
                    i += 1
                    continue
                # 标题处理
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                
                title_text = line[level:].strip()
                if title_text:
                    self.document.add_heading(title_text, level=min(level, 4))
            elif line == '---':
                # 分隔线，可以用段落或其他方式表示
                p = self.document.add_paragraph()
                p.add_run("─" * 50)
            elif line.startswith('**') and line.endswith('**'):
                # 粗体段落，例如章节页码范围
                p = self.document.add_paragraph()
                p.add_run(line.strip('*')).bold = True
            elif '**' in line:
                # 段落中包含粗体文本
                p = self.document.add_paragraph()
                self._parse_mixed_text(p, line)
            else:
                # 普通段落
                if line.strip():
                    self.document.add_paragraph(line)
            
            i += 1

    def _parse_mixed_text(self, paragraph, text):
        """
        解析包含粗体标记的混合文本
        """
        parts = text.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # 普通文本
                if part:
                    paragraph.add_run(part)
            else:
                # 粗体文本
                if part:
                    paragraph.add_run(part).bold = True
